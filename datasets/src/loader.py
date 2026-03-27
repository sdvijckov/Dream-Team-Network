"""
Загрузчик текстовых файлов из различных форматов.
Поддерживает: .txt, .md, .docx
"""

import os
import chardet
from pathlib import Path
from typing import Optional, Dict, List
from docx import Document


class TextLoader:
    """Загрузчик текстов из файлов разных форматов"""
    
    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.docx'}
    
    def __init__(self, encoding: str = 'utf-8'):
        self.default_encoding = encoding
    
    def load_file(self, file_path: str) -> Dict:
        """
        Загружает текст из файла.
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Dict с ключами:
                - text: содержимое файла
                - source_file: имя файла
                - source_path: полный путь
                - extension: расширение файла
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        ext = path.suffix.lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Неподдерживаемое расширение: {ext}. "
                           f"Поддерживаются: {self.SUPPORTED_EXTENSIONS}")
        
        # Загрузка в зависимости от формата
        if ext == '.docx':
            text = self._load_docx(path)
        else:
            text = self._load_text_file(path)
        
        return {
            'text': text,
            'source_file': path.name,
            'source_path': str(path.absolute()),
            'extension': ext
        }
    
    def _load_text_file(self, path: Path) -> str:
        """Загружает .txt или .md файл с автоопределением кодировки"""
        
        # Сначала пробуем стандартную кодировку
        try:
            with open(path, 'r', encoding=self.default_encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            pass
        
        # Если не получилось, определяем кодировку
        with open(path, 'rb') as f:
            raw_data = f.read()
            detected = chardet.detect(raw_data)
            encoding = detected['encoding'] or self.default_encoding
        
        with open(path, 'r', encoding=encoding) as f:
            return f.read()
    
    def _load_docx(self, path: Path) -> str:
        """Извлекает текст из .docx файла"""
        doc = Document(path)
        
        # Извлекаем текст из параграфов
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        return '\n\n'.join(paragraphs)
    
    def load_directory(self, dir_path: str, recursive: bool = True) -> List[Dict]:
        """
        Загружает все текстовые файлы из директории.
        
        Args:
            dir_path: Путь к директории
            recursive: Рекурсивно обходить поддиректории
            
        Returns:
            Список Dict с данными из load_file()
        """
        path = Path(dir_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Директория не найдена: {dir_path}")
        
        files = []
        
        if recursive:
            for ext in self.SUPPORTED_EXTENSIONS:
                files.extend(path.rglob(f'*{ext}'))
        else:
            for ext in self.SUPPORTED_EXTENSIONS:
                files.extend(path.glob(f'*{ext}'))
        
        results = []
        for file_path in sorted(files):
            try:
                data = self.load_file(str(file_path))
                results.append(data)
            except Exception as e:
                print(f"⚠️  Ошибка загрузки {file_path}: {e}")
        
        return results


# Пример использования
if __name__ == '__main__':
    loader = TextLoader()
    
    # Загрузка одного файла
    data = loader.load_file('./raw/books/test.txt')
    print(f"Загружено: {data['source_file']}")
    print(f"Длина текста: {len(data['text'])} символов")
    
    # Загрузка директории
    # files = loader.load_directory('./raw')
    # print(f"Загружено файлов: {len(files)}")
